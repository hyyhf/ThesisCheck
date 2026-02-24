"""Export service for generating review report documents."""

import io
import logging
import re

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor, Cm, Emu
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

from models.schemas import ReviewComment

logger = logging.getLogger(__name__)

# ---------- Font Configuration ----------
CN_FONT = "黑体"
EN_FONT = "Times New Roman"


def _set_run_font(run, font_name_cn=CN_FONT, font_name_en=EN_FONT, size=None, bold=False, color=None, italic=False):
    """Helper to set font properties on a run consistently."""
    run.font.name = font_name_en
    # Ensure rPr and rFonts elements exist
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:ascii="{font_name_en}" w:hAnsi="{font_name_en}" w:eastAsia="{font_name_cn}"/>')
        rPr.insert(0, rFonts)
    else:
        rFonts.set(qn("w:eastAsia"), font_name_cn)
        rFonts.set(qn("w:ascii"), font_name_en)
        rFonts.set(qn("w:hAnsi"), font_name_en)
    if size:
        run.font.size = size
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color


def _set_paragraph_format(para, alignment=None, space_before=None, space_after=None,
                          line_spacing=None, left_indent=None, first_line_indent=None,
                          keep_with_next=False):
    """Helper to set paragraph formatting consistently."""
    fmt = para.paragraph_format
    if alignment is not None:
        para.alignment = alignment
    if space_before is not None:
        fmt.space_before = space_before
    if space_after is not None:
        fmt.space_after = space_after
    if line_spacing is not None:
        fmt.line_spacing = line_spacing
    if left_indent is not None:
        fmt.left_indent = left_indent
    if first_line_indent is not None:
        fmt.first_line_indent = first_line_indent
    if keep_with_next:
        fmt.keep_with_next = True


def _add_styled_heading(doc, text, level=1, color=None):
    """Add a heading paragraph with manual font styling instead of Word built-in styles.

    This ensures the heading uses our designated fonts (SimHei) instead of the
    default theme fonts.

    Level 0: document title (22pt, centered)
    Level 1: section heading (16pt, left, with bottom border)
    Level 2: sub-heading (13pt, left)
    """
    para = doc.add_paragraph()

    if level == 0:
        _set_paragraph_format(
            para,
            alignment=WD_ALIGN_PARAGRAPH.CENTER,
            space_before=Pt(12),
            space_after=Pt(6),
            line_spacing=1.5,
        )
        run = para.add_run(text)
        _set_run_font(run, size=Pt(22), bold=True, color=color or RGBColor(0x1A, 0x1A, 0x2E))

    elif level == 1:
        _set_paragraph_format(
            para,
            space_before=Pt(24),
            space_after=Pt(10),
            line_spacing=1.5,
            keep_with_next=True,
        )
        # Add a bottom border to the paragraph for visual structure
        pPr = para._element.get_or_add_pPr()
        pBdr = parse_xml(
            f'<w:pBdr {nsdecls("w")}>'
            '  <w:bottom w:val="single" w:sz="6" w:space="4" w:color="CCCCCC"/>'
            '</w:pBdr>'
        )
        pPr.append(pBdr)
        run = para.add_run(text)
        _set_run_font(run, size=Pt(16), bold=True, color=color or RGBColor(0x1A, 0x1A, 0x2E))

    elif level == 2:
        _set_paragraph_format(
            para,
            space_before=Pt(14),
            space_after=Pt(6),
            line_spacing=1.4,
            keep_with_next=True,
        )
        run = para.add_run(text)
        _set_run_font(run, size=Pt(13), bold=True, color=color or RGBColor(0x33, 0x33, 0x33))

    return para


def _add_horizontal_rule(doc, color="BBBBBB"):
    """Add a thin horizontal rule using a bordered empty paragraph."""
    para = doc.add_paragraph()
    _set_paragraph_format(para, space_before=Pt(4), space_after=Pt(4))
    pPr = para._element.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:bottom w:val="single" w:sz="4" w:space="1" w:color="{color}"/>'
        '</w:pBdr>'
    )
    pPr.append(pBdr)
    return para


def _add_body_text(doc, text, size=Pt(11), color=None, bold=False, italic=False,
                   left_indent=None, first_line_indent=None, space_after=Pt(6),
                   line_spacing=1.5):
    """Add a body text paragraph with consistent font styling."""
    para = doc.add_paragraph()
    _set_paragraph_format(
        para,
        space_after=space_after,
        line_spacing=line_spacing,
        left_indent=left_indent,
        first_line_indent=first_line_indent,
    )
    run = para.add_run(text)
    _set_run_font(run, size=size, bold=bold, color=color, italic=italic)
    return para


def _add_markdown_paragraph(doc, text, font_size=Pt(12), line_spacing=1.8, space_after=Pt(6),
                            first_line_indent=Pt(24), left_indent=None, color=None):
    """Add a paragraph that parses **bold** markdown markers into Word bold runs.

    Splits text by ** markers and alternates between normal and bold runs.
    All runs use the configured Chinese + English fonts.
    """
    para = doc.add_paragraph()
    _set_paragraph_format(
        para,
        space_after=space_after,
        line_spacing=line_spacing,
        first_line_indent=first_line_indent,
        left_indent=left_indent,
    )

    # Split by **...**  pattern, capturing the bold segments
    parts = re.split(r"\*\*(.+?)\*\*", text)

    for i, part in enumerate(parts):
        if not part:
            continue
        is_bold = (i % 2 == 1)  # odd indices are inside ** **
        run = para.add_run(part)
        _set_run_font(run, size=font_size, bold=is_bold, color=color)

    return para


# Severity display config
SEVERITY_CONFIG = {
    "error": {"label": "严重问题 (必须修改)", "color": RGBColor(0xC6, 0x28, 0x28)},
    "warning": {"label": "一般问题 (建议修改)", "color": RGBColor(0xE6, 0x7E, 0x22)},
    "suggestion": {"label": "改进建议 (可选)", "color": RGBColor(0x27, 0xAE, 0x60)},
}


def _set_default_document_font(doc):
    """Set the default font for the entire document to SimHei."""
    style = doc.styles["Normal"]
    font = style.font
    font.name = EN_FONT
    font.size = Pt(11)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), CN_FONT)

    # Also ensure heading styles use our font
    for heading_level in range(1, 5):
        style_name = f"Heading {heading_level}"
        if style_name in doc.styles:
            h_style = doc.styles[style_name]
            h_font = h_style.font
            h_font.name = EN_FONT
            h_style._element.rPr.rFonts.set(qn("w:eastAsia"), CN_FONT)


def generate_review_report(
    comments: list[ReviewComment],
    summary: str,
    document_title: str,
) -> io.BytesIO:
    """Generate a Word document (.docx) containing the review report.

    Args:
        comments: List of review comments.
        summary: Overall review summary.
        document_title: Title of the original document.

    Returns:
        BytesIO buffer containing the .docx file.
    """
    doc = Document()

    # -- Page margins --
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.8)
        section.right_margin = Cm(2.8)

    # -- Set document default font --
    _set_default_document_font(doc)

    # -- Title --
    _add_styled_heading(doc, "论文评审报告", level=0)

    # -- Decorative line under title --
    _add_horizontal_rule(doc, color="2C3E50")

    # -- Document info --
    _add_body_text(
        doc,
        f"原文档: {document_title}",
        size=Pt(11),
        color=RGBColor(0x55, 0x55, 0x55),
        space_after=Pt(4),
        line_spacing=1.3,
    )

    # -- Statistics --
    error_count = sum(1 for c in comments if c.severity == "error")
    warning_count = sum(1 for c in comments if c.severity == "warning")
    suggestion_count = sum(1 for c in comments if c.severity == "suggestion")

    stats_text = (
        f"共 {len(comments)} 条评审意见: "
        f"{error_count} 个严重问题, "
        f"{warning_count} 个一般问题, "
        f"{suggestion_count} 个改进建议"
    )
    _add_body_text(
        doc,
        stats_text,
        size=Pt(10.5),
        bold=True,
        space_after=Pt(12),
        line_spacing=1.3,
    )

    # -- Summary Section --
    _add_styled_heading(doc, "总体评审摘要", level=1)

    if summary:
        for para_text in summary.split("\n"):
            stripped = para_text.strip()
            if not stripped:
                continue
            _add_markdown_paragraph(
                doc,
                stripped,
                font_size=Pt(11),
                line_spacing=1.6,
                space_after=Pt(8),
                first_line_indent=Pt(22),
            )

    # -- Comments grouped by severity --
    severity_order = ["error", "warning", "suggestion"]

    for severity in severity_order:
        severity_comments = [c for c in comments if c.severity == severity]
        if not severity_comments:
            continue

        config = SEVERITY_CONFIG.get(severity, SEVERITY_CONFIG["suggestion"])
        _add_styled_heading(doc, config["label"], level=1, color=config["color"])

        for idx, comment in enumerate(severity_comments, 1):
            # Comment number sub-heading
            _add_body_text(
                doc,
                f"意见 {idx}",
                size=Pt(12),
                bold=True,
                space_after=Pt(4),
                line_spacing=1.4,
            )

            # Target text (quoted) - indented block with left border effect
            quote_para = doc.add_paragraph()
            _set_paragraph_format(
                quote_para,
                left_indent=Cm(1),
                space_after=Pt(4),
                line_spacing=1.5,
            )
            # Add a left border for quote-style appearance
            pPr = quote_para._element.get_or_add_pPr()
            pBdr = parse_xml(
                f'<w:pBdr {nsdecls("w")}>'
                '  <w:left w:val="single" w:sz="12" w:space="8" w:color="CCCCCC"/>'
                '</w:pBdr>'
            )
            pPr.append(pBdr)

            label_run = quote_para.add_run("原文: ")
            _set_run_font(label_run, size=Pt(10.5), bold=True, color=RGBColor(0x55, 0x55, 0x55))

            text_run = quote_para.add_run(f'"{comment.target_text}"')
            _set_run_font(text_run, size=Pt(10.5), color=RGBColor(0x55, 0x55, 0x55), italic=True)

            # Comment content
            comment_para = doc.add_paragraph()
            _set_paragraph_format(
                comment_para,
                left_indent=Cm(1),
                space_after=Pt(4),
                line_spacing=1.5,
            )
            cl_run = comment_para.add_run("意见: ")
            _set_run_font(cl_run, size=Pt(11), bold=True)

            cc_run = comment_para.add_run(comment.comment)
            _set_run_font(cc_run, size=Pt(11))

            # Paragraph index
            _add_body_text(
                doc,
                f"位置: 第 {comment.paragraph_index + 1} 段",
                size=Pt(9),
                color=RGBColor(0x99, 0x99, 0x99),
                left_indent=Cm(1),
                space_after=Pt(16),
                line_spacing=1.3,
            )

    # Save to buffer
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def generate_comment_report(
    comment_text: str,
    document_title: str,
) -> io.BytesIO:
    """Generate a beautifully formatted Word document (.docx) for overall comment.

    Uses manual font styling (no default Word heading styles) with consistent
    SimHei font. Parses markdown **bold** markers into Word bold formatting.

    Args:
        comment_text: The overall comment text (may contain **bold** markers).
        document_title: Title of the original document.

    Returns:
        BytesIO buffer containing the .docx file.
    """
    doc = Document()

    # Set default document margins
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3)
        section.right_margin = Cm(3)

    # Set default font
    _set_default_document_font(doc)

    # -- Main Title --
    title_para = doc.add_paragraph()
    _set_paragraph_format(
        title_para,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        space_before=Pt(20),
        space_after=Pt(4),
        line_spacing=1.5,
    )
    title_run = title_para.add_run("论文全文评语")
    _set_run_font(title_run, size=Pt(22), bold=True, color=RGBColor(0x1A, 0x1A, 0x2E))

    # -- Subtitle: document name --
    sub_para = doc.add_paragraph()
    _set_paragraph_format(
        sub_para,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=Pt(12),
        line_spacing=1.3,
    )
    sub_run = sub_para.add_run(f"-- {document_title} --")
    _set_run_font(sub_run, size=Pt(14), color=RGBColor(0x55, 0x55, 0x55))

    # -- Divider line --
    _add_horizontal_rule(doc, color="2C3E50")

    # Add some spacing after divider
    spacer = doc.add_paragraph()
    _set_paragraph_format(spacer, space_after=Pt(8))

    # -- Comment body --
    for paragraph_text in comment_text.split("\n"):
        stripped = paragraph_text.strip()
        if not stripped:
            continue
        _add_markdown_paragraph(
            doc,
            stripped,
            font_size=Pt(12),
            line_spacing=1.8,
            space_after=Pt(10),
            first_line_indent=Pt(24),
        )

    # Save to buffer
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
